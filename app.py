import sys
import multiprocessing


from docx import Document
from docx.shared import Inches
import io
import PIL

from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas

from widgets.plot_tab import PlotTab,TabBar
from widgets.menu_widget import MenuWidget
from widgets.plotTab.reel_widget import ReelItem

class MainWindow(QMainWindow):

    def __init__(self):
        super().__init__()
        self.setWindowTitle("MAHLE")
        self.resize(1920, 1080)

        self.tabBar = TabBar(self)
        self.setCentralWidget(self.tabBar)

        #Creating menubar
        self.menubar = MenuWidget(self)
        self.menubar.setGeometry(QRect(0, 0, 800, 27))
        self.menubar.setObjectName("menubar")

        #Creating status bar
        self.statusbar = QStatusBar(self)
        self.statusbar.setObjectName("statusbar")
        self.setStatusBar(self.statusbar)

        self.menubar.fileSelected.connect(self.tabBar.read_csv)
        self.menubar.newWindow.connect(self.create_new_window)
        self.menubar.actionAddCategory.triggered.connect(self.tabBar.addCategory)
        self.menubar.exportReel.triggered.connect(self.writeReel)

        QMetaObject.connectSlotsByName(self)

    def create_new_window(self):
        new_p = multiprocessing.Process(target=run)
        new_p.start()

    def writeReel(self):
        fileName, _ = QFileDialog.getSaveFileName(self.tabBar, "Export Reel", "untitled.docx", "Word Document (*.docx)")
        if not fileName.endswith('.docx'):
            fileName += '.docx'

        doc = Document()

        for tab in self.tabBar.findChildren(PlotTab):
            for item in tab.reel.scroll_content.findChildren(ReelItem, "ReelItem"):
                if item.checkbox.isChecked():
                    img = PIL.Image.fromarray(item.img[:, :, :3])
                    img.resize((400, 300))
                    img_file = io.BytesIO()
                    img.save(img_file, format='PNG')
                    doc.add_picture(img_file, width=Inches(6))

                    doc.add_paragraph(f"Cursor Values:\n\n {item.cursor_vals}")
                    doc.add_paragraph(f"Annotation:\n\n {item.annotation}")
            
        doc.save(fileName)

def run():
    app = QApplication(sys.argv)
    main = MainWindow()
    main.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    multiprocessing.set_start_method('spawn')
    run()